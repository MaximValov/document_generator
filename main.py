import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import os
from docx.shared import Pt, Inches, Cm
from PIL import Image
import tempfile
import base64


def load_substitution_rules(sub_file):
    """Load substitution rules from Excel file (columns: old, new)"""
    try:
        sub_df = pd.read_excel(sub_file, header=None)
        if len(sub_df.columns) >= 2:
            return dict(zip(sub_df[0], sub_df[1]))
        return {}
    except Exception as e:
        st.error(f"Error loading substitution file: {str(e)}")
        return {}


def process_dataframe(df, sub_dict, remove_last_n_rows, remove_cols, round_decimals):
    """Apply all transformations to the dataframe"""
    processed_df = df.copy()
    if remove_last_n_rows and remove_last_n_rows > 0:
        processed_df = processed_df.iloc[:-remove_last_n_rows]
    if remove_cols:
        start_col, end_col = remove_cols
        cols_to_drop = processed_df.columns[start_col - 1:end_col]
        processed_df = processed_df.drop(cols_to_drop, axis=1)
    if sub_dict:
        processed_df.columns = [str(col) for col in processed_df.columns]
        # Only replace exact matches in column names
        processed_df.columns = [sub_dict.get(col, col) for col in processed_df.columns]
        # Only replace exact matches in cell values
        processed_df = processed_df.applymap(lambda x: sub_dict.get(x, x))
    if round_decimals is not None:
        for col in processed_df.select_dtypes(include=['number']).columns:
            processed_df[col] = processed_df[col].apply(
                lambda x: f"{float(x):,.{round_decimals}f}".replace('.', ',')
                if pd.notna(x) else '-'
            )
    processed_df = processed_df.replace([np.nan, 'nan', 'NaN', 'NaT'], '-')
    return processed_df

def add_image_to_cell(cell, image_path, width_cm, height_cm=None, filename=None, show_filename=True):
    """Add an image to a table cell with specified dimensions in cm and filename below"""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    if height_cm:
        run.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))
    else:
        run.add_picture(image_path, width=Cm(width_cm))

    if show_filename and filename:
        paragraph = cell.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run(filename)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)


def create_image_table_doc(image_files, table_rows, table_cols, image_width_cm, table_width_percent, height_cm=None,
                           show_filename=True):
    """Create a Word document with an image table"""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    img_table = doc.add_table(rows=table_rows, cols=table_cols)
    img_table.autofit = False

    # Set table width as percentage
    tbl_pr = img_table._tblPr
    tbl_width = OxmlElement('w:tblW')
    tbl_width.set(qn('w:w'), str(int(table_width_percent * 50)))  # Convert percentage to twentieths of a percent
    tbl_width.set(qn('w:type'), 'pct')
    tbl_pr.append(tbl_width)

    for row in img_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcPr.append(border)

    img_count = 0
    for i in range(table_rows):
        for j in range(table_cols):
            if img_count < len(image_files):
                cell = img_table.cell(i, j)
                cell.paragraphs[0].alignment = 1
                filename = os.path.splitext(image_files[img_count].name)[0]
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                    img = Image.open(image_files[img_count])
                    img.save(tmp.name)
                    tmp_path = tmp.name
                add_image_to_cell(cell, tmp_path, image_width_cm, height_cm, filename, show_filename)
                os.unlink(tmp_path)
                img_count += 1
    return doc


def create_image_table_preview(image_files, table_rows, table_cols, width_cm, height_cm=None, show_filename=True):
    """Create a properly working HTML preview of the image table"""
    preview_container = st.container()
    cols = preview_container.columns(table_cols)
    with preview_container:
        st.markdown("""
        <style>
        .preview-image {
            max-width: 150px;
            max-height: 150px;
            display: block;
            margin: 0 auto;
        }
        .preview-filename {
            font-family: Times New Roman;
            font-size: 10pt;
            text-align: center;
            margin-top: 5px;
            word-break: break-word;
        }
        </style>
        """, unsafe_allow_html=True)

    img_index = 0
    for row in range(table_rows):
        cols = preview_container.columns(table_cols)
        for col in range(table_cols):
            if img_index < len(image_files):
                with cols[col]:
                    filename = os.path.splitext(image_files[img_index].name)[0]
                    img = Image.open(image_files[img_index])
                    st.image(img, use_column_width=True, caption=filename if show_filename else "")
                    img_index += 1
            else:
                with cols[col]:
                    st.write("")


def convert_excel_to_word(df):
    """Convert DataFrame to Word document with borders"""
    doc = Document()
    table = doc.add_table(rows=1, cols=len(df.columns))

    def set_font(cell, text):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def set_cell_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcPr.append(border)

    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        set_font(hdr_cells[i], str(column))
        set_cell_borders(hdr_cells[i])

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            set_font(row_cells[i], str(value))
            set_cell_borders(row_cells[i])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def set_cell_borders(cell):
    """Set cell borders with 0.5pt black borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 4 = 0.5pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcPr.append(border)

def main():
    st.title("Document Generator")
    tab1, tab2, tab3 = st.tabs(["Excel to Word Converter", "Image Table Generator", "Image + Word Table Generator"])
    with tab1:
        st.header("Excel Batch to Word Converter")
        col1, col2 = st.columns(2)
        with col1:
            data_files = st.file_uploader("Upload Excel files to convert",
                                          type=["xlsx", "xls"],
                                          accept_multiple_files=True,
                                          key="excel_uploader")
        with col2:
            sub_file = st.file_uploader("Upload substitution file (sub.xlsx)",
                                        type=["xlsx", "xls"],
                                        help="First column: text to find, Second column: replacement text",
                                        key="sub_uploader")

        if data_files:
            sub_dict = load_substitution_rules(sub_file) if sub_file else {}
            if sub_dict:
                st.info(f"Loaded {len(sub_dict)} substitution rules")
                if st.checkbox("Show substitution rules", key="show_subs"):
                    st.dataframe(pd.DataFrame(list(sub_dict.items()), columns=["Find", "Replace"]))

            with st.expander("Transformation Options", expanded=True):
                cols = st.columns(3)
                with cols[0]:
                    remove_rows = st.checkbox("Remove last N rows", key="remove_rows")
                    if remove_rows:
                        n_rows = st.number_input("Number of rows to remove from end", 1, 100, 1, key="n_rows")
                with cols[1]:
                    remove_cols = st.checkbox("Remove columns", key="remove_cols")
                    if remove_cols:
                        col_range = st.slider("Column range to remove", 1, 50, (1, 1), key="col_range")
                with cols[2]:
                    round_enabled = st.checkbox("Round numbers", key="round_enabled")
                    round_decimals = st.number_input("Decimal places", 0, 6, 2,
                                                     key="decimals") if round_enabled else None

            for data_file in data_files:
                try:
                    original_df = pd.read_excel(data_file)
                    file_name = os.path.splitext(data_file.name)[0]
                    with st.expander(f"Processing: {file_name}", expanded=True):
                        processed_df = process_dataframe(
                            original_df,
                            sub_dict,
                            n_rows if remove_rows else None,
                            col_range if remove_cols else None,
                            round_decimals if round_enabled else None
                        )
                        st.subheader("Complete Modified Table Preview")
                        st.dataframe(processed_df, height=400)
                        word_buffer = convert_excel_to_word(processed_df)
                        st.download_button(
                            label=f"Download {file_name}.docx",
                            data=word_buffer,
                            file_name=f"{file_name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_{file_name}"
                        )
                except Exception as e:
                    st.error(f"Error processing {data_file.name}: {str(e)}")
    with tab2:
        st.header("Image Table Generator")
        image_files = st.file_uploader("Upload images for the table",
                                       type=["png", "jpg", "jpeg", "bmp"],
                                       accept_multiple_files=True,
                                       key="image_uploader")

        if image_files:
            with st.expander("Image Table Configuration", expanded=True):
                cols = st.columns(3)
                with cols[0]:
                    table_rows = st.number_input("Table rows", 1, 20, 1, key="img_rows")
                with cols[1]:
                    table_cols = st.number_input("Table columns", 1, 10, min(3, len(image_files)), key="img_cols")
                with cols[2]:
                    table_width_percent = st.number_input("Table width (%)", 1, 100, 100, 1, key="table_width_percent")

                cols = st.columns(2)
                with cols[0]:
                    image_width_cm = st.number_input("Image width (cm)", 0.5, 30.0, 5.0, 0.1, key="img_width_cm")
                with cols[1]:
                    fixed_height = st.checkbox("Fixed height", key="fixed_height")
                    if fixed_height:
                        height_cm = st.number_input("Image height (cm)", 0.5, 30.0, 5.0, 0.1, key="img_height_cm")
                    else:
                        height_cm = None
                show_filename = st.checkbox("Show filename", value=True, key="show_filename")

            if st.button("Preview Image Table", key="preview_img_table"):
                with st.spinner("Generating preview..."):
                    try:
                        st.subheader("Table Preview")
                        create_image_table_preview(
                            image_files,
                            table_rows,
                            table_cols,
                            image_width_cm,
                            height_cm,
                            show_filename
                        )
                        st.info(
                            "Note: This is an approximation of how the table will look in Word.")
                    except Exception as e:
                        st.error(f"Error generating preview: {str(e)}")

            if st.button("Generate Image + Table Document", key="generate_img_table_tab3"):
                with st.spinner("Creating document..."):
                    try:
                        doc = Document()
                        style = doc.styles['Normal']
                        font = style.font
                        font.name = 'Times New Roman'
                        font.size = Pt(12)
        
                        # Create one continuous table with 2 columns
                        table = doc.add_table(rows=len(image_files), cols=2)
                        table.autofit = False
        
                        # Set table width to 100% of page
                        tbl_pr = table._tblPr
                        tbl_width = OxmlElement('w:tblW')
                        tbl_width.set(qn('w:w'), "5000")  # 5000 = 100% width
                        tbl_width.set(qn('w:type'), 'pct')
                        tbl_pr.append(tbl_width)
        
                        # Set table borders to 0.5pt
                        tbl_borders = OxmlElement('w:tblBorders')
                        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set(qn('w:val'), 'single')
                            border.set(qn('w:sz'), '4')  # 4 = 0.5pt
                            border.set(qn('w:space'), '0')
                            border.set(qn('w:color'), '000000')
                            tbl_borders.append(border)
                        tbl_pr.append(tbl_borders)
        
                        for idx, img_file in enumerate(image_files):
                            img_name = os.path.splitext(img_file.name)[0]
                            row = table.rows[idx]
        
                            # Set row height to 0.6cm
                            tr_pr = row._tr.get_or_add_trPr()
                            tr_height = OxmlElement('w:trHeight')
                            tr_height.set(qn('w:val'), str(int(0.6 * 567)))  # 0.6cm in twentieths of a point
                            tr_height.set(qn('w:hRule'), 'exact')
                            tr_pr.append(tr_height)
        
                            # Left cell (image)
                            left_cell = row.cells[0]
                            left_cell.width = Cm(image_width_cm)
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                                img = Image.open(img_file)
                                img.save(tmp.name)
                                add_image_to_cell(left_cell, tmp.name, image_width_cm,
                                                  show_filename=show_filename, filename=img_name)
                                os.unlink(tmp.name)
        
                            # Right cell (table)
                            right_cell = row.cells[1]
                            right_cell.width = Cm(table_width_cm)
        
                            if img_name in table_mapping:
                                try:
                                    src_doc = Document(table_mapping[img_name])
                                    if src_doc.tables:
                                        src_table = src_doc.tables[0]
        
                                        # Create nested table in right cell
                                        nested_table = right_cell.add_table(
                                            rows=len(src_table.rows),
                                            cols=len(src_table.columns))
        
                                        # Copy content and formatting
                                        for i, src_row in enumerate(src_table.rows):
                                            for j, src_cell in enumerate(src_row.cells):
                                                new_cell = nested_table.cell(i, j)
                                                # Clear existing paragraphs
                                                for para in new_cell.paragraphs:
                                                    p = para._element
                                                    p.getparent().remove(p)
                                                # Copy content
                                                for para in src_cell.paragraphs:
                                                    new_para = new_cell.add_paragraph()
                                                    for run in para.runs:
                                                        new_run = new_para.add_run(run.text)
                                                        new_run.bold = run.bold
                                                        new_run.italic = run.italic
                                                        new_run.underline = run.underline
                                                        new_run.font.name = run.font.name or 'Times New Roman'
                                                        new_run.font.size = run.font.size or Pt(12)
                                                # Set cell borders
                                                set_cell_borders(new_cell)
                                except Exception as e:
                                    right_cell.text = f"Error loading table: {str(e)}"
        
                        # Save the document
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
        
                        st.success("Document created successfully!")
                        st.download_button(
                            label="Download Word Document",
                            data=buffer,
                            file_name="images_and_tables.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_img_table_tab3"
                        )
        
                    except Exception as e:
                        st.error(f"Error creating document: {str(e)}")
                
    with tab3:
        st.header("Image + Word Table Generator")
        image_files = st.file_uploader("Upload images for the left column",
                                       type=["png", "jpg", "jpeg", "bmp"],
                                       accept_multiple_files=True,
                                       key="image_table_uploader")

        table_files = st.file_uploader("Upload Word files with tables for the right column (same names as images)",
                                       type=["docx"],
                                       accept_multiple_files=True,
                                       key="table_uploader")

        if image_files and table_files:
            # Create a mapping of filename (without extension) to file objects
            table_mapping = {os.path.splitext(f.name)[0]: f for f in table_files}

            with st.expander("Configuration", expanded=True):
                cols = st.columns(2)
                with cols[0]:
                    image_width_cm = st.number_input("Image width (cm)", 0.5, 30.0, 5.0, 0.1, key="img_width_cm_tab3")
                with cols[1]:
                    table_width_cm = st.number_input("Table width (cm)", 0.5, 30.0, 15.0, 0.1,
                                                     key="table_width_cm_tab3")

                show_filename = st.checkbox("Show filename", value=False, key="show_filename_tab3")

            if st.button("Preview Image + Table", key="preview_img_table_tab3"):
                st.subheader("Preview")
                for img_file in image_files[:5]:  # Limit preview to 5 items
                    img_name = os.path.splitext(img_file.name)[0]
                    if img_name in table_mapping:
                        col1, col2 = st.columns(2)
                        with col1:
                            st.image(img_file, caption=img_name if show_filename else "", width=200)
                        with col2:
                            try:
                                doc = Document(table_mapping[img_name])
                                style = doc.styles['Normal']
                                font = style.font
                                font.name = 'Times New Roman'
                                font.size = Pt(12)
                                for table in doc.tables:
                                    html = []
                                    for row in table.rows:
                                        html.append("<tr>")
                                        for cell in row.cells:
                                            html.append(f"<td>{cell.text}</td>")
                                        html.append("</tr>")
                                    st.markdown(f"<table>{''.join(html)}</table>", unsafe_allow_html=True)
                                    break  # Just show first table
                            except Exception as e:
                                st.error(f"Error loading table for {img_name}: {str(e)}")
                    else:
                        st.warning(f"No matching table found for image: {img_name}")


            if st.button("Generate Image + Table Document", key="generate_img_table_tab3"):
                with st.spinner("Creating document..."):
                    try:
                        doc = Document()
                        style = doc.styles['Normal']
                        font = style.font
                        font.name = 'Times New Roman'
                        font.size = Pt(12)

                        for img_file in image_files:
                            img_name = os.path.splitext(img_file.name)[0]
                            if img_name in table_mapping:
                                # Add a new table with 2 columns
                                table = doc.add_table(rows=1, cols=2)
                                table.autofit = False

                                # Set table width to 100% of page (fixed)
                                tbl_pr = table._tblPr
                                tbl_width = OxmlElement('w:tblW')
                                tbl_width.set(qn('w:w'), "5000")  # 5000 = 100% width in twentieths of a percent
                                tbl_width.set(qn('w:type'), 'pct')
                                tbl_pr.append(tbl_width)

                                # Set column widths
                                cols = table.columns
                                cols[0].width = Cm(image_width_cm)
                                cols[1].width = Cm(table_width_cm)

                                # Set table borders to 0.5pt
                                tbl_borders = OxmlElement('w:tblBorders')
                                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                                    border = OxmlElement(f'w:{border_name}')
                                    border.set(qn('w:val'), 'single')
                                    border.set(qn('w:sz'), '4')  # 4 = 0.5pt
                                    border.set(qn('w:space'), '0')
                                    border.set(qn('w:color'), '000000')
                                    tbl_borders.append(border)
                                tbl_pr.append(tbl_borders)

                                # Add image to left cell
                                left_cell = table.cell(0, 0)
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                                    img = Image.open(img_file)
                                    img.save(tmp.name)
                                    add_image_to_cell(left_cell, tmp.name, image_width_cm,
                                                      show_filename=show_filename, filename=img_name)
                                    os.unlink(tmp.name)

                                # Add Word table to right cell
                                right_cell = table.cell(0, 1)
                                try:
                                    # Open the source Word document
                                    src_doc = Document(table_mapping[img_name])

                                    # Get the first table from source document
                                    if src_doc.tables:
                                        src_table = src_doc.tables[0]

                                        # Create new table in right cell
                                        new_table = right_cell.add_table(
                                            rows=len(src_table.rows),
                                            cols=len(src_table.columns)
                                        )

                                        # Set table borders to 0.5pt
                                        new_tbl_pr = new_table._tblPr
                                        new_tbl_borders = OxmlElement('w:tblBorders')
                                        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                                            border = OxmlElement(f'w:{border_name}')
                                            border.set(qn('w:val'), 'single')
                                            border.set(qn('w:sz'), '4')  # 4 = 0.5pt
                                            border.set(qn('w:space'), '0')
                                            border.set(qn('w:color'), '000000')
                                            new_tbl_borders.append(border)
                                        new_tbl_pr.append(new_tbl_borders)

                                        # Copy content and formatting
                                        for i, row in enumerate(src_table.rows):
                                            for j, cell in enumerate(row.cells):
                                                new_cell = new_table.cell(i, j)
                                                # Clear existing paragraphs to avoid duplication
                                                for para in new_cell.paragraphs:
                                                    p = para._element
                                                    p.getparent().remove(p)

                                                # Copy text and formatting
                                                for para in cell.paragraphs:
                                                    new_para = new_cell.add_paragraph()
                                                    for run in para.runs:
                                                        new_run = new_para.add_run(run.text)
                                                        new_run.bold = run.bold
                                                        new_run.italic = run.italic
                                                        new_run.underline = run.underline
                                                        new_run.font.name = run.font.name or 'Times New Roman'
                                                        new_run.font.size = run.font.size or Pt(12)

                                                # Set cell borders to 0.5pt
                                                set_cell_borders(new_cell)

                                except Exception as e:
                                    right_cell.text = f"Error loading table: {str(e)}"

                                # Add space between items
                                doc.add_paragraph()

                        # Save the document
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)

                        st.success("Document created successfully!")
                        st.download_button(
                            label="Download Word Document",
                            data=buffer,
                            file_name="images_and_tables.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_img_table_tab3"
                        )

                    except Exception as e:
                        st.error(f"Error creating document: {str(e)}")

if __name__ == "__main__":
    st.set_page_config(layout="wide")
    main()
